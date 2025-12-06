#!/usr/bin/env python3
"""
Flask API for YouTube Transcript Downloader

Endpoints:
  POST /api/single - Download single video transcript
  POST /api/batch - Upload CSV and download transcripts in batch
  POST /api/playlist - Download all transcripts from a playlist
  POST /api/channel - Download all transcripts from a channel
  GET /api/download/<filename> - Download generated files

Run: python3 flask_transcript_api.py
Test with Postman on: http://localhost:5000
"""

import argparse
import contextlib
import csv
import itertools
import json
import os
import re
import shutil
import subprocess
import sys
import time
import zipfile
import random
import traceback
from datetime import timedelta
from typing import Any, Dict, List, Optional
from flask import Flask, request, jsonify, send_file
from werkzeug.utils import secure_filename
import tempfile

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()
app.config['OUTPUT_FOLDER'] = os.path.join(tempfile.gettempdir(), 'transcript_outputs')

os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# ----------------- Utilities -----------------

def safe_filename(name: str) -> str:
    name = (name or "").strip()
    name = name.replace("/", "_").replace("\\", "_")
    name = re.sub(r"[^A-Za-z0-9._\-\s]", "_", name)
    return name[:200]

def extract_video_id(url: str) -> str:
    from urllib.parse import urlparse, parse_qs
    parsed = urlparse(url)
    if parsed.query:
        q = parse_qs(parsed.query)
        if "v" in q and q["v"]:
            return q["v"][0]
    path = (parsed.path or "").strip("/")
    parts = path.split("/")
    if parts:
        candidate = parts[-1]
        if re.match(r"^[A-Za-z0-9_-]{6,}$", candidate):
            return candidate
    raise ValueError(f"Could not extract video id from URL: {url}")

# ----------------- Proxy helpers -----------------

@contextlib.contextmanager
def use_proxy_env(proxy_url: Optional[str]):
    old_http = os.environ.get("HTTP_PROXY")
    old_https = os.environ.get("HTTPS_PROXY")
    if not proxy_url:
        try:
            yield
            return
        finally:
            return
    if not proxy_url.startswith("http://") and not proxy_url.startswith("https://"):
        proxy_val = "http://" + proxy_url
    else:
        proxy_val = proxy_url
    os.environ["HTTP_PROXY"] = proxy_val
    os.environ["HTTPS_PROXY"] = proxy_val
    try:
        yield
    finally:
        if old_http is not None:
            os.environ["HTTP_PROXY"] = old_http
        else:
            os.environ.pop("HTTP_PROXY", None)
        if old_https is not None:
            os.environ["HTTPS_PROXY"] = old_https
        else:
            os.environ.pop("HTTPS_PROXY", None)

def load_proxies_from_file(path: str) -> List[str]:
    proxies = []
    with open(path, "r", encoding="utf-8") as fh:
        for ln in fh:
            ln = ln.strip()
            if not ln or ln.startswith("#"):
                continue
            proxies.append(ln)
    return proxies

def proxy_round_robin(proxies: List[str]):
    if not proxies:
        while True:
            yield None
    for p in itertools.cycle(proxies):
        yield p

# ----------------- Backoff -----------------

def maybe_sleep_backoff(attempt: int, base: float = 1.0, cap: float = 30.0):
    wait = min(cap, base * (2 ** attempt))
    jitter = random.uniform(-0.5 * wait, 0.5 * wait)
    sleep_for = max(0.1, wait + jitter)
    print(f"Backoff: sleeping for {sleep_for:.1f}s (attempt {attempt})")
    time.sleep(sleep_for)

# ----------------- Transcript retrieval -----------------

def normalize_transcript_obj(obj: Any) -> List[Dict[str, Any]]:
    def to_item_dict(it):
        d = {}
        if isinstance(it, dict):
            d["text"] = it.get("text")
            d["start"] = it.get("start")
            d["duration"] = it.get("duration")
            d["speaker"] = it.get("speaker", None)
            d["confidence"] = it.get("confidence", None)
        else:
            d["text"] = getattr(it, "text", None)
            d["start"] = getattr(it, "start", None)
            d["duration"] = getattr(it, "duration", None)
            d["speaker"] = getattr(it, "speaker", None)
            d["confidence"] = getattr(it, "confidence", None)
        d.setdefault("text", None)
        d.setdefault("start", None)
        d.setdefault("duration", None)
        d.setdefault("speaker", None)
        d.setdefault("confidence", None)
        return d

    if isinstance(obj, (list, tuple)):
        return [to_item_dict(it) for it in obj]
    try:
        items_iter = list(obj)
        if items_iter:
            return [to_item_dict(it) for it in items_iter]
    except Exception:
        pass
    for method_name in ("fetch", "list", "to_list", "as_list", "get_transcript"):
        if hasattr(obj, method_name):
            try:
                res = getattr(obj, method_name)
                res = res() if callable(res) else res
                return normalize_transcript_obj(res)
            except Exception:
                pass
    raise RuntimeError("Unable to normalize transcript object returned by library.")

def yt_dlp_subtitle_fallback(url: str, lang: str = "en") -> Optional[List[Dict[str, Any]]]:
    try:
        from yt_dlp import YoutubeDL
    except Exception:
        YoutubeDL = None

    vid = extract_video_id(url)
    out_base = f"{vid}.subs"
    if YoutubeDL:
        ydl_opts = {
            'skip_download': True,
            'writesubtitles': True,
            'writeautomaticsub': True,
            'subtitleslangs': [lang],
            'subtitlesformat': 'vtt/srt/best',
            'outtmpl': out_base + '.%(ext)s',
            'quiet': True,
            'no_warnings': True,
        }
        try:
            with YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
        except Exception:
            pass
    else:
        cmd = [
            "yt-dlp",
            "--write-auto-sub",
            "--sub-lang", lang,
            "--skip-download",
            "-o", out_base + ".%(ext)s",
            url,
        ]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except Exception:
            return None

    candidates = [f for f in os.listdir(".") if f.startswith(out_base) and f.endswith((".vtt", ".srt", ".txt"))]
    if not candidates:
        return None

    candidate = candidates[0]
    with open(candidate, "r", encoding="utf-8", errors="ignore") as fh:
        data = fh.read()

    cues = []
    vtt_pattern = re.compile(r"(\d{2}:\d{2}:\d{2}\.\d+)\s*-->\s*(\d{2}:\d{2}:\d{2}\.\d+)\s*\n(.*?)(?=\n{2,}|\Z)", re.S)
    for mm in vtt_pattern.finditer(data):
        start_ts = mm.group(1).strip()
        end_ts = mm.group(2).strip()
        text = mm.group(3).strip().replace("\n", " ")
        def to_secs(t):
            h, m, s = t.split(":")
            return int(h) * 3600 + int(m) * 60 + float(s)
        try:
            start = to_secs(start_ts)
            duration = round(to_secs(end_ts) - start, 3)
        except Exception:
            start = None
            duration = None
        cues.append({"text": text, "start": start, "duration": duration, "speaker": None, "confidence": None})
    if cues:
        return cues

    srt_pattern = re.compile(r"\d+\s*\n(\d{2}:\d{2}:\d{2},\d+)\s*-->\s*(\d{2}:\d{2}:\d{2},\d+)\s*\n(.*?)(?=\n{2,}|\Z)", re.S)
    for mm in srt_pattern.finditer(data):
        start_ts = mm.group(1).strip().replace(",", ".")
        end_ts = mm.group(2).strip().replace(",", ".")
        text = mm.group(3).strip().replace("\n", " ")
        def to_secs(t):
            h, m, s = t.split(":")
            return int(h) * 3600 + int(m) * 60 + float(s)
        try:
            start = to_secs(start_ts)
            duration = round(to_secs(end_ts) - start, 3)
        except Exception:
            start = None
            duration = None
        cues.append({"text": text, "start": start, "duration": duration, "speaker": None, "confidence": None})
    if cues:
        return cues

    parts = [p.strip() for p in re.split(r"\n{2,}", data) if p.strip()]
    for p in parts:
        p_clean = re.sub(r"\d+\n", "", p)
        p_clean = p_clean.replace("\n", " ").strip()
        if p_clean:
            cues.append({"text": p_clean, "start": None, "duration": None, "speaker": None, "confidence": None})
    return cues if cues else None

def download_transcript(url: str, prefer_lang: Optional[str] = None, max_retries: int = 4) -> List[Dict[str, Any]]:
    last_error = None
    vid = extract_video_id(url)

    for attempt in range(max_retries):
        try:
            try:
                from youtube_transcript_api import YouTubeTranscriptApi
            except Exception as e:
                YouTubeTranscriptApi = None
                last_error = e
            
            items = None
            if YouTubeTranscriptApi:
                try:
                    if hasattr(YouTubeTranscriptApi, "get_transcript"):
                        raw = YouTubeTranscriptApi.get_transcript(vid, languages=[prefer_lang] if prefer_lang else None)
                    else:
                        api = YouTubeTranscriptApi()
                        try:
                            raw = api.fetch(vid, languages=[prefer_lang] if prefer_lang else None)
                        except TypeError:
                            raw = api.fetch(vid)
                    items = normalize_transcript_obj(raw)
                except Exception as e:
                    last_error = e
                    msg = str(e).lower()
                    if "429" in msg or "too many requests" in msg:
                        print("Received 429; will retry with backoff.")
                        maybe_sleep_backoff(attempt)
                        continue
                    if ("transcript" in msg and "disabled" in msg) or ("no transcript" in msg):
                        print("Transcript unavailable via API:", e)
                        items = None
                        break
                    if "http error" in msg or "request" in msg or "blocked" in msg:
                        print("Network-like error; retrying...")
                        maybe_sleep_backoff(attempt)
                        continue
                    break

            if items:
                normalized = []
                for it in items:
                    normalized.append({
                        "text": it.get("text") if isinstance(it, dict) else getattr(it, "text", None),
                        "start": it.get("start") if isinstance(it, dict) else getattr(it, "start", None),
                        "duration": it.get("duration") if isinstance(it, dict) else getattr(it, "duration", None),
                        "speaker": it.get("speaker", None) if isinstance(it, dict) else getattr(it, "speaker", None),
                        "confidence": it.get("confidence", None) if isinstance(it, dict) else getattr(it, "confidence", None),
                    })
                return normalized

            try:
                subs = yt_dlp_subtitle_fallback(url, lang=prefer_lang or "en")
                if subs:
                    return subs
            except Exception as e:
                last_error = e

            break

        except Exception as e:
            last_error = e
            msg = str(e).lower()
            if "429" in msg or "too many requests" in msg:
                maybe_sleep_backoff(attempt)
                continue
            if attempt < max_retries - 1:
                maybe_sleep_backoff(attempt)
                continue
            else:
                break

    if last_error:
        raise last_error
    raise RuntimeError("Failed to download transcript")

# ----------------- Format writers -----------------

def secs_to_srt_timestamp(secs: float) -> str:
    td = timedelta(seconds=secs)
    total_seconds = int(td.total_seconds())
    hh = total_seconds // 3600
    mm = (total_seconds % 3600) // 60
    ss = total_seconds % 60
    ms = int((td.total_seconds() - total_seconds) * 1000)
    return f"{hh:02d}:{mm:02d}:{ss:02d},{ms:03d}"

def secs_to_vtt_timestamp(secs: float) -> str:
    td = timedelta(seconds=secs)
    total_seconds = int(td.total_seconds())
    hh = total_seconds // 3600
    mm = (total_seconds % 3600) // 60
    ss = total_seconds % 60
    ms = int((td.total_seconds() - total_seconds) * 1000)
    return f"{hh:02d}:{mm:02d}:{ss:02d}.{ms:03d}"

def compute_end(start: Optional[float], duration: Optional[float], default_next_start: Optional[float]) -> Optional[float]:
    if start is None:
        return None
    if duration is not None:
        return start + duration
    if default_next_start is not None:
        return default_next_start
    return None

def write_srt(items: List[Dict[str, Any]], path: str) -> str:
    lines = []
    n = len(items)
    for i, it in enumerate(items):
        start = it.get("start")
        duration = it.get("duration")
        next_start = None
        if i + 1 < n:
            next_start = items[i+1].get("start")
        end = compute_end(start, duration, next_start)
        if start is None:
            start = (i * 2.0)
            end = start + (duration if duration else 2.0)
        if end is None:
            end = start + (duration if duration else 2.0)
        idx = i + 1
        lines.append(str(idx))
        lines.append(f"{secs_to_srt_timestamp(start)} --> {secs_to_srt_timestamp(end)}")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        if speaker:
            lines.append(f"{speaker}: {text}")
        else:
            lines.append(text)
        lines.append("")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return path

def write_vtt(items: List[Dict[str, Any]], path: str) -> str:
    lines = ["WEBVTT", ""]
    n = len(items)
    for i, it in enumerate(items):
        start = it.get("start")
        duration = it.get("duration")
        next_start = None
        if i + 1 < n:
            next_start = items[i+1].get("start")
        end = compute_end(start, duration, next_start)
        if start is None:
            start = (i * 2.0)
            end = start + (duration if duration else 2.0)
        if end is None:
            end = start + (duration if duration else 2.0)
        lines.append(f"{secs_to_vtt_timestamp(start)} --> {secs_to_vtt_timestamp(end)}")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        if speaker:
            lines.append(f"{speaker}: {text}")
        else:
            lines.append(text)
        lines.append("")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return path

def write_csv(items: List[Dict[str, Any]], path: str) -> str:
    fieldnames = ["start", "duration", "speaker", "confidence", "text"]
    def fmt(ts):
        if ts is None:
            return ""
        return secs_to_vtt_timestamp(ts)
    with open(path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for it in items:
            writer.writerow({
                "start": fmt(it.get("start")),
                "duration": fmt(it.get("duration")),
                "speaker": it.get("speaker"),
                "confidence": it.get("confidence"),
                "text": it.get("text"),
            })
    return path

def write_txt(items: List[Dict[str, Any]], path: str) -> str:
    with open(path, "w", encoding="utf-8") as fh:
        for it in items:
            speaker = it.get("speaker")
            text = it.get("text") or ""
            if speaker:
                fh.write(f"[{speaker}] {text}\n")
            else:
                fh.write(f"{text}\n")
    return path

def write_json(items: List[Dict[str, Any]], path: str) -> str:
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(items, fh, ensure_ascii=False, indent=2)
    return path

def write_docx(items: List[Dict[str, Any]], path: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:
        raise RuntimeError("python-docx is required for docx output. pip install python-docx") from e

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    for it in items:
        start = it.get("start")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        ts = secs_to_vtt_timestamp(start) if start is not None else ""
        para = doc.add_paragraph()
        if speaker:
            para.add_run(f"{speaker} ").bold = True
        if ts:
            para.add_run(f"[{ts}] ").italic = True
        para.add_run(text)
    doc.save(path)
    return path

def write_pdf(items: List[Dict[str, Any]], path: str) -> str:
    try:
        from fpdf import FPDF
    except Exception as e:
        raise RuntimeError("fpdf is required for pdf output. pip install fpdf") from e

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)

    for it in items:
        start = it.get("start")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        ts = secs_to_vtt_timestamp(start) if start is not None else ""
        header = ""
        if speaker:
            header = f"{speaker} "
        if ts:
            header += f"[{ts}] "
        if header:
            pdf.set_font("Arial", style='B', size=11)
            pdf.multi_cell(0, 6, header)
            pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 6, text)
        pdf.ln(2)
    pdf.output(path)
    return path

def save_output(vid: str, items: List[Dict[str, Any]], fmt: str, custom_path: Optional[str]) -> str:
    if custom_path:
        out_path = custom_path
    else:
        out_path = f"{vid}_transcript.{fmt}"
    fmt = fmt.lower()
    if fmt == "txt":
        return write_txt(items, out_path)
    if fmt == "json":
        return write_json(items, out_path)
    if fmt == "srt":
        return write_srt(items, out_path)
    if fmt == "vtt":
        return write_vtt(items, out_path)
    if fmt == "csv":
        return write_csv(items, out_path)
    if fmt == "docx":
        return write_docx(items, out_path)
    if fmt == "pdf":
        return write_pdf(items, out_path)
    raise ValueError(f"Unsupported format: {fmt}")

# ----------------- Playlist & Channel helpers -----------------

def extract_playlist_videos(playlist_url: str) -> List[Dict[str, str]]:
    try:
        from yt_dlp import YoutubeDL
    except Exception:
        YoutubeDL = None
    if YoutubeDL:
        ydl_opts = {"extract_flat": "in_playlist", "skip_download": True, "quiet": True}
        with YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(playlist_url, download=False)
        entries = info.get("entries", []) or []
        out = []
        for e in entries:
            vid = e.get("id")
            title = e.get("title") or e.get("fulltitle") or ""
            if vid:
                out.append({"id": vid, "title": title, "url": f"https://www.youtube.com/watch?v={vid}"})
        return out
    else:
        try:
            p = subprocess.run(["yt-dlp", "-J", "--flat-playlist", playlist_url], capture_output=True, text=True, check=True)
            info = json.loads(p.stdout)
            entries = info.get("entries", []) or []
            out = []
            for e in entries:
                vid = e.get("id")
                title = e.get("title") or ""
                if vid:
                    out.append({"id": vid, "title": title, "url": f"https://www.youtube.com/watch?v={vid}"})
            return out
        except Exception as e:
            raise RuntimeError(f"yt-dlp required for playlist extraction: {e}")

def apply_output_template(template: Optional[str], index: int, video_id: str, title: str, fmt: str) -> str:
    safe_title = safe_filename(title) if title else ""
    if not template:
        return f"{video_id}_transcript.{fmt}"
    name = template.replace("{index}", str(index)).replace("{video_id}", video_id).replace("{title}", safe_title).replace("{ext}", fmt)
    if not name.lower().endswith("." + fmt):
        name = name + "." + fmt
    return safe_filename(name)

def get_channel_id_from_url(channel_url: str) -> str:
    try:
        from yt_dlp import YoutubeDL
    except Exception:
        YoutubeDL = None
    if YoutubeDL:
        ydl_opts = {"quiet": True, "skip_download": True, "extract_flat": True}
        with YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(channel_url, download=False)
            cid = info.get("channel_id") or info.get("uploader_id") or info.get("id")
            if cid:
                return cid
    try:
        p = subprocess.run(["yt-dlp", "-J", "--no-warnings", channel_url], capture_output=True, text=True, check=True)
        info = json.loads(p.stdout)
        cid = info.get("channel_id") or info.get("uploader_id") or info.get("id")
        if cid:
            return cid
    except Exception as e:
        raise RuntimeError(f"Failed to resolve channel id for {channel_url!r}: {e}")
    raise RuntimeError(f"Could not determine channel id for {channel_url!r}")

def channel_uploads_playlist_url(channel_id: str) -> str:
    if not isinstance(channel_id, str) or not channel_id.startswith("UC") or len(channel_id) <= 2:
        raise ValueError("Invalid channel id (expected string starting with 'UC').")
    uploads_id = "UU" + channel_id[2:]
    return f"https://www.youtube.com/playlist?list={uploads_id}"

def process_row(url: str, fmt: str, outputfile: Optional[str], outdir: str, lang: Optional[str]) -> Dict[str, Any]:
    try:
        fmt = (fmt or "json").lower()
        if fmt not in ("txt","json","srt","vtt","docx","pdf","csv"):
            return {"ok": False, "error": f"unsupported format: {fmt}", "url": url, "outfile": outputfile}
        items = download_transcript(url, prefer_lang=lang)
        vid = extract_video_id(url)
        if outputfile:
            fname = safe_filename(outputfile)
        else:
            fname = f"{vid}_transcript.{fmt}"
        if not fname.lower().endswith("." + fmt):
            fname = fname + "." + fmt
        target_path = os.path.join(outdir, fname)
        saved = save_output(vid, items, fmt, target_path)
        return {"ok": True, "path": saved, "url": url}
    except Exception as e:
        return {"ok": False, "error": str(e), "url": url, "outfile": outputfile}

# ----------------- Flask Routes -----------------

@app.route('/')
def index():
    return jsonify({
        "message": "YouTube Transcript Downloader API",
        "endpoints": {
            "POST /api/single": "Download single video transcript",
            "POST /api/batch": "Upload CSV and batch download",
            "POST /api/playlist": "Download playlist transcripts",
            "POST /api/channel": "Download channel transcripts",
            "GET /api/download/<filename>": "Download generated files"
        }
    })

@app.route('/api/single', methods=['POST'])
def single_transcript():
    """
    Body params:
    - url (required): YouTube video URL
    - format: txt, json, srt, vtt, csv, docx, pdf (default: json)
    - lang: language code (default: en)
    - output_filename: custom output filename
    """
    try:
        data = request.get_json() if request.is_json else request.form.to_dict()
        
        url = data.get('url')
        if not url:
            return jsonify({"error": "URL is required"}), 400
        
        fmt = data.get('format', 'json').lower()
        lang = data.get('lang', 'en')
        output_filename = data.get('output_filename')
        
        if fmt not in ("txt", "json", "srt", "vtt", "docx", "pdf", "csv"):
            return jsonify({"error": f"Unsupported format: {fmt}"}), 400
        
        # Download transcript
        items = download_transcript(url, prefer_lang=lang)
        vid = extract_video_id(url)
        
        # Save output
        if output_filename:
            fname = safe_filename(output_filename)
        else:
            fname = f"{vid}_transcript.{fmt}"
        
        if not fname.lower().endswith("." + fmt):
            fname = fname + "." + fmt
        
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], fname)
        saved_path = save_output(vid, items, fmt, output_path)
        
        return jsonify({
            "success": True,
            "video_id": vid,
            "format": fmt,
            "filename": fname,
            "download_url": f"/api/download/{fname}"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/batch', methods=['POST'])
def batch_transcript():
    """
    Upload CSV and batch download
    Form-data: file (CSV with columns: URL, format, outputfileName)
    Query params: ?lang=en&create_zip=true
    """
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        lang = request.args.get('lang', 'en')
        create_zip = request.args.get('create_zip', 'true').lower() == 'true'
        
        # Save CSV
        filename = secure_filename(file.filename)
        csv_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(csv_path)
        
        # Create output directory
        batch_id = f"batch_{int(time.time())}"
        outdir = os.path.join(app.config['OUTPUT_FOLDER'], batch_id)
        os.makedirs(outdir, exist_ok=True)
        
        # Process batch
        results = []
        with open(csv_path, newline="", encoding="utf-8") as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                url = row.get("URL") or row.get("url") or ""
                fmt = row.get("format") or row.get("Format") or "json"
                outfn = row.get("outputfileName") or row.get("outputFileName") or None
                
                res = process_row(url, fmt, outfn, outdir, lang)
                results.append(res)
        
        # Create zip if requested
        zip_filename = None
        if create_zip:
            zip_filename = f"{batch_id}.zip"
            zip_path = os.path.join(app.config['OUTPUT_FOLDER'], zip_filename)
            with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                for f in os.listdir(outdir):
                    zf.write(os.path.join(outdir, f), arcname=f)
            shutil.rmtree(outdir, ignore_errors=True)
        
        ok_count = sum(1 for r in results if r.get("ok"))
        err_count = sum(1 for r in results if not r.get("ok"))
        
        response = {
            "success": True,
            "batch_id": batch_id,
            "total": len(results),
            "successful": ok_count,
            "failed": err_count,
            "results": results
        }
        
        if zip_filename:
            response["download_url"] = f"/api/download/{zip_filename}"
            response["zip_file"] = zip_filename
        
        return jsonify(response)
    
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/playlist', methods=['POST'])
def playlist_transcript():
    """
    Download all transcripts from a playlist
    Body (JSON or form-data):
    - url (required): YouTube playlist URL
    - format: output format (default: json)
    - lang: language code (default: en)
    - output_template: filename template with {index}, {video_id}, {title}, {ext}
    - create_zip: true/false (default: true)
    """
    try:
        data = request.get_json() if request.is_json else request.form.to_dict()
        
        url = data.get('url')
        if not url:
            return jsonify({"error": "URL is required"}), 400
        
        fmt = data.get('format', 'json').lower()
        lang = data.get('lang', 'en')
        output_template = data.get('output_template')
        create_zip = str(data.get('create_zip', 'true')).lower() == 'true'
        
        if fmt not in ("txt", "json", "srt", "vtt", "docx", "pdf", "csv"):
            return jsonify({"error": f"Unsupported format: {fmt}"}), 400
        
        # Extract videos
        videos = extract_playlist_videos(url)
        if not videos:
            return jsonify({"error": "No videos found in playlist"}), 404
        
        # Create output directory
        playlist_id = f"playlist_{int(time.time())}"
        outdir = os.path.join(app.config['OUTPUT_FOLDER'], playlist_id)
        os.makedirs(outdir, exist_ok=True)
        
        # Process videos
        results = []
        for idx, v in enumerate(videos, start=1):
            video_url = v.get("url")
            vid = v.get("id")
            title = v.get("title") or ""
            outname = apply_output_template(output_template, idx, vid, title, fmt)
            
            res = process_row(video_url, fmt, outname, outdir, lang)
            results.append({**res, "title": title, "index": idx})
        
        # Create zip if requested
        zip_filename = None
        if create_zip:
            zip_filename = f"{playlist_id}.zip"
            zip_path = os.path.join(app.config['OUTPUT_FOLDER'], zip_filename)
            with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                for f in os.listdir(outdir):
                    zf.write(os.path.join(outdir, f), arcname=f)
            shutil.rmtree(outdir, ignore_errors=True)
        
        ok_count = sum(1 for r in results if r.get("ok"))
        err_count = sum(1 for r in results if not r.get("ok"))
        
        response = {
            "success": True,
            "playlist_id": playlist_id,
            "total_videos": len(videos),
            "successful": ok_count,
            "failed": err_count,
            "results": results
        }
        
        if zip_filename:
            response["download_url"] = f"/api/download/{zip_filename}"
            response["zip_file"] = zip_filename
        
        return jsonify(response)
    
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/channel', methods=['POST'])
def channel_transcript():
    """
    Download all transcripts from a channel
    Body (JSON or form-data):
    - url (required): YouTube channel URL
    - format: output format (default: json)
    - lang: language code (default: en)
    - output_template: filename template with {index}, {video_id}, {title}, {ext}
    - create_zip: true/false (default: true)
    """
    try:
        data = request.get_json() if request.is_json else request.form.to_dict()
        
        url = data.get('url')
        if not url:
            return jsonify({"error": "URL is required"}), 400
        
        fmt = data.get('format', 'json').lower()
        lang = data.get('lang', 'en')
        output_template = data.get('output_template')
        create_zip = str(data.get('create_zip', 'true')).lower() == 'true'
        
        if fmt not in ("txt", "json", "srt", "vtt", "docx", "pdf", "csv"):
            return jsonify({"error": f"Unsupported format: {fmt}"}), 400
        
        # Get channel ID and uploads playlist
        channel_id = get_channel_id_from_url(url)
        playlist_url = channel_uploads_playlist_url(channel_id)
        
        # Extract videos
        videos = extract_playlist_videos(playlist_url)
        if not videos:
            return jsonify({"error": "No videos found in channel"}), 404
        
        # Create output directory
        channel_dir_id = f"channel_{int(time.time())}"
        outdir = os.path.join(app.config['OUTPUT_FOLDER'], channel_dir_id)
        os.makedirs(outdir, exist_ok=True)
        
        # Process videos
        results = []
        for idx, v in enumerate(videos, start=1):
            video_url = v.get("url")
            vid = v.get("id")
            title = v.get("title") or ""
            outname = apply_output_template(output_template, idx, vid, title, fmt)
            
            res = process_row(video_url, fmt, outname, outdir, lang)
            results.append({**res, "title": title, "index": idx})
        
        # Create zip if requested
        zip_filename = None
        if create_zip:
            zip_filename = f"{channel_dir_id}.zip"
            zip_path = os.path.join(app.config['OUTPUT_FOLDER'], zip_filename)
            with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                for f in os.listdir(outdir):
                    zf.write(os.path.join(outdir, f), arcname=f)
            shutil.rmtree(outdir, ignore_errors=True)
        
        ok_count = sum(1 for r in results if r.get("ok"))
        err_count = sum(1 for r in results if not r.get("ok"))
        
        response = {
            "success": True,
            "channel_id": channel_id,
            "total_videos": len(videos),
            "successful": ok_count,
            "failed": err_count,
            "results": results
        }
        
        if zip_filename:
            response["download_url"] = f"/api/download/{zip_filename}"
            response["zip_file"] = zip_filename
        
        return jsonify(response)
    
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download generated files"""
    try:
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
        
        return send_file(file_path, as_attachment=True, download_name=filename)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    print("=" * 60)
    print("YouTube Transcript Downloader API")
    print("=" * 60)
    print("Server running on: http://localhost:5000")
    print("\nAvailable endpoints:")
    print("  POST /api/single     - Download single video")
    print("  POST /api/batch      - Batch download from CSV")
    print("  POST /api/playlist   - Download playlist")
    print("  POST /api/channel    - Download channel")
    print("  GET  /api/download/<filename> - Download file")
    print("\nPress CTRL+C to stop")
    print("=" * 60)
    app.run(debug=True, host='0.0.0.0', port=5000)