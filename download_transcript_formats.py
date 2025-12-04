#!/usr/bin/env python3
"""
download_transcript_formats.py

Usage:
  python3 download_transcript_formats.py "YOUTUBE_URL" --format json --output out.json

Supported formats: txt (default), json, srt, vtt, docx, pdf, csv
"""

import argparse
import csv
import json
import os
import re
import subprocess
from datetime import timedelta
from typing import Any, Dict, List, Optional
from urllib.parse import urlparse, parse_qs

# --- Utilities: extract video id ------------------------------------------------
def extract_video_id(url: str) -> str:
    parsed = urlparse(url)
    if parsed.query:
        q = parse_qs(parsed.query)
        if "v" in q and q["v"]:
            return q["v"][0]
    path = (parsed.path or "").strip("/")
    parts = path.split("/")
    if parts:
        candidate = parts[-1]
        if re.match(r"^[A-Za-z0-9_-]{6,}$", candidate):
            return candidate
    raise ValueError(f"Could not extract video id from URL: {url}")

# --- Normalize transcript object so every item has text,start,duration,speaker,confidence
def normalize_transcript_obj(obj: Any) -> List[Dict[str, Any]]:
    def to_item_dict(it):
        d = {}
        if isinstance(it, dict):
            d["text"] = it.get("text")
            d["start"] = it.get("start")
            d["duration"] = it.get("duration")
            # keep existing speaker/confidence if any
            d["speaker"] = it.get("speaker", None)
            d["confidence"] = it.get("confidence", None)
        else:
            d["text"] = getattr(it, "text", None)
            d["start"] = getattr(it, "start", None)
            d["duration"] = getattr(it, "duration", None)
            d["speaker"] = getattr(it, "speaker", None)
            d["confidence"] = getattr(it, "confidence", None)
        # ensure keys exist
        d.setdefault("text", None)
        d.setdefault("start", None)
        d.setdefault("duration", None)
        d.setdefault("speaker", None)
        d.setdefault("confidence", None)
        return d

    # If list/tuple:
    if isinstance(obj, (list, tuple)):
        return [to_item_dict(it) for it in obj]

    # Iterable
    try:
        items_iter = list(obj)
        if items_iter:
            return [to_item_dict(it) for it in items_iter]
    except Exception:
        pass

    # Methods to obtain list
    for method_name in ("fetch", "list", "to_list", "as_list", "get_transcript"):
        if hasattr(obj, method_name):
            try:
                res = getattr(obj, method_name)
                res = res() if callable(res) else res
                return normalize_transcript_obj(res)
            except Exception:
                pass

    raise RuntimeError("Unable to normalize transcript object returned by library.")

# --- Transcription retrieval (youtube-transcript-api with yt-dlp fallback) -----
def yt_dlp_subtitle_fallback(url: str, lang: str = "en") -> Optional[List[Dict[str, Any]]]:
    """
    Try using yt-dlp to download auto subtitles and return list of dicts with text,start(optional),duration(optional)
    """
    try:
        from yt_dlp import YoutubeDL
    except Exception:
        YoutubeDL = None

    vid = extract_video_id(url)
    out_base = f"{vid}.subs"

    if YoutubeDL:
        ydl_opts = {
            'skip_download': True,
            'writesubtitles': True,
            'writeautomaticsub': True,
            'subtitleslangs': [lang],
            'subtitlesformat': 'vtt/srt/best',
            'outtmpl': out_base + '.%(ext)s',
            'quiet': True,
            'no_warnings': True,
        }
        try:
            with YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
        except Exception:
            pass
    else:
        cmd = [
            "yt-dlp",
            "--write-auto-sub",
            "--sub-lang", lang,
            "--skip-download",
            "-o", out_base + ".%(ext)s",
            url,
        ]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except Exception:
            return None

    # find candidates
    candidates = [f for f in os.listdir(".") if f.startswith(out_base) and f.endswith((".vtt", ".srt", ".txt"))]
    if not candidates:
        return None

    candidate = candidates[0]
    with open(candidate, "r", encoding="utf-8", errors="ignore") as fh:
        data = fh.read()

    cues = []

    # VTT
    vtt_pattern = re.compile(r"(\d{2}:\d{2}:\d{2}\.\d+)\s*-->\s*(\d{2}:\d{2}:\d{2}\.\d+)\s*\n(.*?)(?=\n{2,}|\Z)", re.S)
    for mm in vtt_pattern.finditer(data):
        start_ts = mm.group(1).strip()
        end_ts = mm.group(2).strip()
        text = mm.group(3).strip().replace("\n", " ")
        def to_secs(t):
            h, m, s = t.split(":")
            return int(h) * 3600 + int(m) * 60 + float(s)
        try:
            start = to_secs(start_ts)
            duration = round(to_secs(end_ts) - start, 3)
        except Exception:
            start = None
            duration = None
        cues.append({"text": text, "start": start, "duration": duration, "speaker": None, "confidence": None})

    if cues:
        return cues

    # SRT
    srt_pattern = re.compile(r"\d+\s*\n(\d{2}:\d{2}:\d{2},\d+)\s*-->\s*(\d{2}:\d{2}:\d{2},\d+)\s*\n(.*?)(?=\n{2,}|\Z)", re.S)
    for mm in srt_pattern.finditer(data):
        start_ts = mm.group(1).strip().replace(",", ".")
        end_ts = mm.group(2).strip().replace(",", ".")
        text = mm.group(3).strip().replace("\n", " ")
        def to_secs(t):
            h, m, s = t.split(":")
            return int(h) * 3600 + int(m) * 60 + float(s)
        try:
            start = to_secs(start_ts)
            duration = round(to_secs(end_ts) - start, 3)
        except Exception:
            start = None
            duration = None
        cues.append({"text": text, "start": start, "duration": duration, "speaker": None, "confidence": None})

    if cues:
        return cues

    # fallback: split by blank lines
    parts = [p.strip() for p in re.split(r"\n{2,}", data) if p.strip()]
    for p in parts:
        p_clean = re.sub(r"\d+\n", "", p)
        p_clean = p_clean.replace("\n", " ").strip()
        if p_clean:
            cues.append({"text": p_clean, "start": None, "duration": None, "speaker": None, "confidence": None})

    return cues if cues else None

def download_transcript(url: str, prefer_lang: Optional[str]) -> List[Dict[str, Any]]:
    """
    Attempts retrieval using youtube-transcript-api first, then yt-dlp fallback.
    Returns normalized list of items with text/start/duration/speaker/confidence
    """
    items = None
    fetch_error = None
    try:
        import youtube_transcript_api
        from youtube_transcript_api import YouTubeTranscriptApi
    except Exception as e:
        YouTubeTranscriptApi = None
        fetch_error = e

    if YouTubeTranscriptApi is not None:
        try:
            # old static API?
            if hasattr(YouTubeTranscriptApi, "get_transcript"):
                raw = YouTubeTranscriptApi.get_transcript(extract_video_id(url), languages=[prefer_lang] if prefer_lang else None)
                items = normalize_transcript_obj(raw)
            else:
                api = YouTubeTranscriptApi()
                try:
                    raw = api.fetch(extract_video_id(url), languages=[prefer_lang] if prefer_lang else None)
                except TypeError:
                    raw = api.fetch(extract_video_id(url))
                items = normalize_transcript_obj(raw)
        except Exception as e:
            fetch_error = e
            items = None

    if not items:
        # fallback to yt-dlp
        items = yt_dlp_subtitle_fallback(url, lang=prefer_lang or "en")

    if not items:
        raise RuntimeError(f"Failed to retrieve transcript. Last error: {fetch_error}")

    # ensure normalized and ensure speaker/confidence keys exist
    normalized = []
    for it in items:
        normalized.append({
            "text": it.get("text") if isinstance(it, dict) else getattr(it, "text", None),
            "start": it.get("start") if isinstance(it, dict) else getattr(it, "start", None),
            "duration": it.get("duration") if isinstance(it, dict) else getattr(it, "duration", None),
            "speaker": it.get("speaker", None) if isinstance(it, dict) else getattr(it, "speaker", None),
            "confidence": it.get("confidence", None) if isinstance(it, dict) else getattr(it, "confidence", None),
        })
    return normalized

# --- Format writers ------------------------------------------------------------
def secs_to_srt_timestamp(secs: float) -> str:
    # returns "HH:MM:SS,mmm"
    td = timedelta(seconds=secs)
    total_seconds = int(td.total_seconds())
    hh = total_seconds // 3600
    mm = (total_seconds % 3600) // 60
    ss = total_seconds % 60
    ms = int((td.total_seconds() - total_seconds) * 1000)
    return f"{hh:02d}:{mm:02d}:{ss:02d},{ms:03d}"

def secs_to_vtt_timestamp(secs: float) -> str:
    # returns "HH:MM:SS.mmm"
    td = timedelta(seconds=secs)
    total_seconds = int(td.total_seconds())
    hh = total_seconds // 3600
    mm = (total_seconds % 3600) // 60
    ss = total_seconds % 60
    ms = int((td.total_seconds() - total_seconds) * 1000)
    return f"{hh:02d}:{mm:02d}:{ss:02d}.{ms:03d}"

def compute_end(start: Optional[float], duration: Optional[float], default_next_start: Optional[float]) -> Optional[float]:
    if start is None:
        return None
    if duration is not None:
        return start + duration
    if default_next_start is not None:
        return default_next_start
    return None

def write_srt(items: List[Dict[str, Any]], path: str) -> str:
    # For end times, if duration missing use next start - small epsilon
    lines = []
    n = len(items)
    for i, it in enumerate(items):
        start = it.get("start")
        duration = it.get("duration")
        next_start = None
        if i + 1 < n:
            next_start = items[i+1].get("start")
        end = compute_end(start, duration, next_start)
        if start is None:
            # if no timing, put 0 timestamps incrementally
            start = (i * 2.0)
            end = start + (duration if duration else 2.0)
        if end is None:
            end = start + (duration if duration else 2.0)
        idx = i + 1
        lines.append(str(idx))
        lines.append(f"{secs_to_srt_timestamp(start)} --> {secs_to_srt_timestamp(end)}")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        if speaker:
            lines.append(f"{speaker}: {text}")
        else:
            lines.append(text)
        lines.append("")  # blank
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return path

def write_vtt(items: List[Dict[str, Any]], path: str) -> str:
    lines = ["WEBVTT", ""]
    n = len(items)
    for i, it in enumerate(items):
        start = it.get("start")
        duration = it.get("duration")
        next_start = None
        if i + 1 < n:
            next_start = items[i+1].get("start")
        end = compute_end(start, duration, next_start)
        if start is None:
            start = (i * 2.0)
            end = start + (duration if duration else 2.0)
        if end is None:
            end = start + (duration if duration else 2.0)
        lines.append(f"{secs_to_vtt_timestamp(start)} --> {secs_to_vtt_timestamp(end)}")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        if speaker:
            lines.append(f"{speaker}: {text}")
        else:
            lines.append(text)
        lines.append("")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return path

def write_csv(items: List[Dict[str, Any]], path: str) -> str:
    fieldnames = ["start", "duration", "speaker", "confidence", "text"]

    def fmt(ts):
        if ts is None:
            return ""
        return secs_to_vtt_timestamp(ts)  # SAME as PDF/DOCX/VTT

    with open(path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()

        for it in items:
            writer.writerow({
                "start": fmt(it.get("start")),                # formatted timestamp
                "duration": fmt(it.get("duration")),          # also formatted (optional)
                "speaker": it.get("speaker"),
                "confidence": it.get("confidence"),
                "text": it.get("text"),
            })

    return path

def write_txt(items: List[Dict[str, Any]], path: str) -> str:
    with open(path, "w", encoding="utf-8") as fh:
        for it in items:
            speaker = it.get("speaker")
            text = it.get("text") or ""
            if speaker:
                fh.write(f"[{speaker}] {text}\n")
            else:
                fh.write(f"{text}\n")
    return path

def write_json(items: List[Dict[str, Any]], path: str) -> str:
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(items, fh, ensure_ascii=False, indent=2)
    return path

def write_docx(items: List[Dict[str, Any]], path: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:
        raise RuntimeError("python-docx is required for docx output. pip install python-docx") from e

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    for it in items:
        start = it.get("start")
        duration = it.get("duration")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        if start is not None:
            ts = f"{secs_to_vtt_timestamp(start)}"
        else:
            ts = ""
        para = doc.add_paragraph()
        if speaker:
            para.add_run(f"{speaker} ").bold = True
        if ts:
            para.add_run(f"[{ts}] ").italic = True
        para.add_run(text)
    doc.save(path)
    return path

def write_pdf(items: List[Dict[str, Any]], path: str) -> str:
    # Use fpdf to create a simple PDF. pip install fpdf
    try:
        from fpdf import FPDF
    except Exception as e:
        raise RuntimeError("fpdf is required for pdf output. pip install fpdf") from e

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)

    for it in items:
        start = it.get("start")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        ts = secs_to_vtt_timestamp(start) if start is not None else ""
        header = ""
        if speaker:
            header = f"{speaker} "
        if ts:
            header += f"[{ts}] "
        # write header bold-like by using same font but adding newline
        if header:
            pdf.set_font("Arial", style='B', size=11)
            pdf.multi_cell(0, 6, header)
            pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 6, text)
        pdf.ln(2)
    pdf.output(path)
    return path

# --- Save output dispatch ------------------------------------------------------
def save_output(vid: str, items: List[Dict[str, Any]], fmt: str, custom_path: Optional[str]) -> str:
    # choose default filename if custom not provided
    if custom_path:
        out_path = custom_path
    else:
        out_path = f"{vid}_transcript.{fmt}"

    fmt = fmt.lower()
    if fmt == "txt":
        return write_txt(items, out_path)
    if fmt == "json":
        return write_json(items, out_path)
    if fmt == "srt":
        return write_srt(items, out_path)
    if fmt == "vtt":
        return write_vtt(items, out_path)
    if fmt == "csv":
        return write_csv(items, out_path)
    if fmt == "docx":
        return write_docx(items, out_path)
    if fmt == "pdf":
        return write_pdf(items, out_path)
    raise ValueError(f"Unsupported format: {fmt}")

# --- CLI ----------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Download YouTube transcript in multiple formats")
    parser.add_argument("url", help="YouTube video URL")
    parser.add_argument("--lang", help="Language code (e.g. en)", default=None)
    parser.add_argument("--format", "-f", choices=("txt","json","srt","vtt","docx","pdf","csv"), default="json", help="output format")
    parser.add_argument("--output", "-o", help="Custom output filename", default=None)
    args = parser.parse_args()

    print("Fetching transcript...")
    items = download_transcript(args.url, prefer_lang=args.lang)
    vid = extract_video_id(args.url)

    # ensure every item has required keys and consistent types
    for it in items:
        it.setdefault("text", None)
        it.setdefault("start", None)
        it.setdefault("duration", None)
        it.setdefault("speaker", None)
        it.setdefault("confidence", None)

    print("Saving as", args.format, "...")
    out = save_output(vid, items, args.format, args.output)
    print("Saved:", out)

if __name__ == "__main__":
    main()

