#!/usr/bin/env python3
"""
download_transcript_formats.py

Single and batch modes:
 - Single mode (existing): python3 download_transcript_formats.py "YOUTUBE_URL" --format json --output out.json
 - Batch mode: python3 download_transcript_formats.py --batch input.csv --outdir outputs --zip

Batch CSV columns (header):
 URL,format,outputfileName
 - format is one of: txt,json,srt,vtt,docx,pdf,csv
 - outputfileName optional: if empty, defaults to VIDEOID_transcript.<ext>

Outputs are saved under --outdir (default ./downloads). If --zip is set, a zip archive of the directory is created.
"""

import argparse
import csv
import json
import os
import re
import subprocess
import sys
import zipfile
from datetime import timedelta
from typing import Any, Dict, List, Optional
from urllib.parse import urlparse, parse_qs

# --------------------- Utilities & transcript retrieval ------------------------

def extract_video_id(url: str) -> str:
    parsed = urlparse(url)
    if parsed.query:
        q = parse_qs(parsed.query)
        if "v" in q and q["v"]:
            return q["v"][0]
    path = (parsed.path or "").strip("/")
    parts = path.split("/")
    if parts:
        candidate = parts[-1]
        if re.match(r"^[A-Za-z0-9_-]{6,}$", candidate):
            return candidate
    raise ValueError(f"Could not extract video id from URL: {url}")

def normalize_transcript_obj(obj: Any) -> List[Dict[str, Any]]:
    def to_item_dict(it):
        d = {}
        if isinstance(it, dict):
            d["text"] = it.get("text")
            d["start"] = it.get("start")
            d["duration"] = it.get("duration")
            d["speaker"] = it.get("speaker", None)
            d["confidence"] = it.get("confidence", None)
        else:
            d["text"] = getattr(it, "text", None)
            d["start"] = getattr(it, "start", None)
            d["duration"] = getattr(it, "duration", None)
            d["speaker"] = getattr(it, "speaker", None)
            d["confidence"] = getattr(it, "confidence", None)
        d.setdefault("text", None)
        d.setdefault("start", None)
        d.setdefault("duration", None)
        d.setdefault("speaker", None)
        d.setdefault("confidence", None)
        return d

    if isinstance(obj, (list, tuple)):
        return [to_item_dict(it) for it in obj]
    try:
        items_iter = list(obj)
        if items_iter:
            return [to_item_dict(it) for it in items_iter]
    except Exception:
        pass
    for method_name in ("fetch", "list", "to_list", "as_list", "get_transcript"):
        if hasattr(obj, method_name):
            try:
                res = getattr(obj, method_name)
                res = res() if callable(res) else res
                return normalize_transcript_obj(res)
            except Exception:
                pass
    raise RuntimeError("Unable to normalize transcript object returned by library.")

def yt_dlp_subtitle_fallback(url: str, lang: str = "en") -> Optional[List[Dict[str, Any]]]:
    try:
        from yt_dlp import YoutubeDL
    except Exception:
        YoutubeDL = None

    vid = extract_video_id(url)
    out_base = f"{vid}.subs"
    if YoutubeDL:
        ydl_opts = {
            'skip_download': True,
            'writesubtitles': True,
            'writeautomaticsub': True,
            'subtitleslangs': [lang],
            'subtitlesformat': 'vtt/srt/best',
            'outtmpl': out_base + '.%(ext)s',
            'quiet': True,
            'no_warnings': True,
        }
        try:
            with YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
        except Exception:
            pass
    else:
        cmd = [
            "yt-dlp",
            "--write-auto-sub",
            "--sub-lang", lang,
            "--skip-download",
            "-o", out_base + ".%(ext)s",
            url,
        ]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except Exception:
            return None

    candidates = [f for f in os.listdir(".") if f.startswith(out_base) and f.endswith((".vtt", ".srt", ".txt"))]
    if not candidates:
        return None

    candidate = candidates[0]
    with open(candidate, "r", encoding="utf-8", errors="ignore") as fh:
        data = fh.read()

    cues = []
    vtt_pattern = re.compile(r"(\d{2}:\d{2}:\d{2}\.\d+)\s*-->\s*(\d{2}:\d{2}:\d{2}\.\d+)\s*\n(.*?)(?=\n{2,}|\Z)", re.S)
    for mm in vtt_pattern.finditer(data):
        start_ts = mm.group(1).strip()
        end_ts = mm.group(2).strip()
        text = mm.group(3).strip().replace("\n", " ")
        def to_secs(t):
            h, m, s = t.split(":")
            return int(h) * 3600 + int(m) * 60 + float(s)
        try:
            start = to_secs(start_ts)
            duration = round(to_secs(end_ts) - start, 3)
        except Exception:
            start = None
            duration = None
        cues.append({"text": text, "start": start, "duration": duration, "speaker": None, "confidence": None})
    if cues:
        return cues

    srt_pattern = re.compile(r"\d+\s*\n(\d{2}:\d{2}:\d{2},\d+)\s*-->\s*(\d{2}:\d{2}:\d{2},\d+)\s*\n(.*?)(?=\n{2,}|\Z)", re.S)
    for mm in srt_pattern.finditer(data):
        start_ts = mm.group(1).strip().replace(",", ".")
        end_ts = mm.group(2).strip().replace(",", ".")
        text = mm.group(3).strip().replace("\n", " ")
        def to_secs(t):
            h, m, s = t.split(":")
            return int(h) * 3600 + int(m) * 60 + float(s)
        try:
            start = to_secs(start_ts)
            duration = round(to_secs(end_ts) - start, 3)
        except Exception:
            start = None
            duration = None
        cues.append({"text": text, "start": start, "duration": duration, "speaker": None, "confidence": None})
    if cues:
        return cues

    parts = [p.strip() for p in re.split(r"\n{2,}", data) if p.strip()]
    for p in parts:
        p_clean = re.sub(r"\d+\n", "", p)
        p_clean = p_clean.replace("\n", " ").strip()
        if p_clean:
            cues.append({"text": p_clean, "start": None, "duration": None, "speaker": None, "confidence": None})
    return cues if cues else None

def download_transcript(url: str, prefer_lang: Optional[str]) -> List[Dict[str, Any]]:
    items = None
    fetch_error = None
    try:
        import youtube_transcript_api
        from youtube_transcript_api import YouTubeTranscriptApi
    except Exception as e:
        YouTubeTranscriptApi = None
        fetch_error = e

    if YouTubeTranscriptApi is not None:
        try:
            vid = extract_video_id(url)
            if hasattr(YouTubeTranscriptApi, "get_transcript"):
                raw = YouTubeTranscriptApi.get_transcript(vid, languages=[prefer_lang] if prefer_lang else None)
                items = normalize_transcript_obj(raw)
            else:
                api = YouTubeTranscriptApi()
                try:
                    raw = api.fetch(vid, languages=[prefer_lang] if prefer_lang else None)
                except TypeError:
                    raw = api.fetch(vid)
                items = normalize_transcript_obj(raw)
        except Exception as e:
            fetch_error = e
            items = None

    if not items:
        items = yt_dlp_subtitle_fallback(url, lang=prefer_lang or "en")

    if not items:
        raise RuntimeError(f"Failed to retrieve transcript. Last error: {fetch_error}")

    normalized = []
    for it in items:
        normalized.append({
            "text": it.get("text") if isinstance(it, dict) else getattr(it, "text", None),
            "start": it.get("start") if isinstance(it, dict) else getattr(it, "start", None),
            "duration": it.get("duration") if isinstance(it, dict) else getattr(it, "duration", None),
            "speaker": it.get("speaker", None) if isinstance(it, dict) else getattr(it, "speaker", None),
            "confidence": it.get("confidence", None) if isinstance(it, dict) else getattr(it, "confidence", None),
        })
    return normalized

# ----------------------- Format writers (same as before) -----------------------

from datetime import timedelta

def secs_to_srt_timestamp(secs: float) -> str:
    td = timedelta(seconds=secs)
    total_seconds = int(td.total_seconds())
    hh = total_seconds // 3600
    mm = (total_seconds % 3600) // 60
    ss = total_seconds % 60
    ms = int((td.total_seconds() - total_seconds) * 1000)
    return f"{hh:02d}:{mm:02d}:{ss:02d},{ms:03d}"

def secs_to_vtt_timestamp(secs: float) -> str:
    td = timedelta(seconds=secs)
    total_seconds = int(td.total_seconds())
    hh = total_seconds // 3600
    mm = (total_seconds % 3600) // 60
    ss = total_seconds % 60
    ms = int((td.total_seconds() - total_seconds) * 1000)
    return f"{hh:02d}:{mm:02d}:{ss:02d}.{ms:03d}"

def compute_end(start: Optional[float], duration: Optional[float], default_next_start: Optional[float]) -> Optional[float]:
    if start is None:
        return None
    if duration is not None:
        return start + duration
    if default_next_start is not None:
        return default_next_start
    return None

def write_srt(items: List[Dict[str, Any]], path: str) -> str:
    lines = []
    n = len(items)
    for i, it in enumerate(items):
        start = it.get("start")
        duration = it.get("duration")
        next_start = None
        if i + 1 < n:
            next_start = items[i+1].get("start")
        end = compute_end(start, duration, next_start)
        if start is None:
            start = (i * 2.0)
            end = start + (duration if duration else 2.0)
        if end is None:
            end = start + (duration if duration else 2.0)
        idx = i + 1
        lines.append(str(idx))
        lines.append(f"{secs_to_srt_timestamp(start)} --> {secs_to_srt_timestamp(end)}")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        if speaker:
            lines.append(f"{speaker}: {text}")
        else:
            lines.append(text)
        lines.append("")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return path

def write_vtt(items: List[Dict[str, Any]], path: str) -> str:
    lines = ["WEBVTT", ""]
    n = len(items)
    for i, it in enumerate(items):
        start = it.get("start")
        duration = it.get("duration")
        next_start = None
        if i + 1 < n:
            next_start = items[i+1].get("start")
        end = compute_end(start, duration, next_start)
        if start is None:
            start = (i * 2.0)
            end = start + (duration if duration else 2.0)
        if end is None:
            end = start + (duration if duration else 2.0)
        lines.append(f"{secs_to_vtt_timestamp(start)} --> {secs_to_vtt_timestamp(end)}")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        if speaker:
            lines.append(f"{speaker}: {text}")
        else:
            lines.append(text)
        lines.append("")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return path

def write_csv(items: List[Dict[str, Any]], path: str) -> str:
    fieldnames = ["start", "duration", "speaker", "confidence", "text"]

    def fmt(ts):
        if ts is None:
            return ""
        return secs_to_vtt_timestamp(ts)

    with open(path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for it in items:
            writer.writerow({
                "start": fmt(it.get("start")),
                "duration": fmt(it.get("duration")),
                "speaker": it.get("speaker"),
                "confidence": it.get("confidence"),
                "text": it.get("text"),
            })
    return path

def write_txt(items: List[Dict[str, Any]], path: str) -> str:
    with open(path, "w", encoding="utf-8") as fh:
        for it in items:
            speaker = it.get("speaker")
            text = it.get("text") or ""
            if speaker:
                fh.write(f"[{speaker}] {text}\n")
            else:
                fh.write(f"{text}\n")
    return path

def write_json(items: List[Dict[str, Any]], path: str) -> str:
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(items, fh, ensure_ascii=False, indent=2)
    return path

def write_docx(items: List[Dict[str, Any]], path: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:
        raise RuntimeError("python-docx is required for docx output. pip install python-docx") from e

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    for it in items:
        start = it.get("start")
        duration = it.get("duration")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        if start is not None:
            ts = f"{secs_to_vtt_timestamp(start)}"
        else:
            ts = ""
        para = doc.add_paragraph()
        if speaker:
            para.add_run(f"{speaker} ").bold = True
        if ts:
            para.add_run(f"[{ts}] ").italic = True
        para.add_run(text)
    doc.save(path)
    return path

def write_pdf(items: List[Dict[str, Any]], path: str) -> str:
    try:
        from fpdf import FPDF
    except Exception as e:
        raise RuntimeError("fpdf is required for pdf output. pip install fpdf") from e

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)

    for it in items:
        start = it.get("start")
        speaker = it.get("speaker")
        text = it.get("text") or ""
        ts = secs_to_vtt_timestamp(start) if start is not None else ""
        header = ""
        if speaker:
            header = f"{speaker} "
        if ts:
            header += f"[{ts}] "
        if header:
            pdf.set_font("Arial", style='B', size=11)
            pdf.multi_cell(0, 6, header)
            pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 6, text)
        pdf.ln(2)
    pdf.output(path)
    return path

def save_output(vid: str, items: List[Dict[str, Any]], fmt: str, custom_path: Optional[str]) -> str:
    if custom_path:
        out_path = custom_path
    else:
        out_path = f"{vid}_transcript.{fmt}"

    fmt = fmt.lower()
    if fmt == "txt":
        return write_txt(items, out_path)
    if fmt == "json":
        return write_json(items, out_path)
    if fmt == "srt":
        return write_srt(items, out_path)
    if fmt == "vtt":
        return write_vtt(items, out_path)
    if fmt == "csv":
        return write_csv(items, out_path)
    if fmt == "docx":
        return write_docx(items, out_path)
    if fmt == "pdf":
        return write_pdf(items, out_path)
    raise ValueError(f"Unsupported format: {fmt}")

# ----------------------- Batch processing -------------------------------------

def safe_filename(name: str) -> str:
    # basic sanitization
    name = name.strip()
    # remove path separators
    name = name.replace("/", "_").replace("\\", "_")
    # remove illegal chars
    name = re.sub(r"[^A-Za-z0-9._\-\s]", "_", name)
    return name

def process_row(url: str, fmt: str, outputfile: Optional[str], outdir: str, lang: Optional[str]) -> Dict[str, Any]:
    """
    Downloads transcript for one row and saves it to outdir.
    Returns a dict with status and message.
    """
    try:
        fmt = (fmt or "json").lower()
        if fmt not in ("txt","json","srt","vtt","docx","pdf","csv"):
            return {"ok": False, "error": f"unsupported format: {fmt}", "url": url, "outfile": outputfile}
        items = download_transcript(url, prefer_lang=lang)
        vid = extract_video_id(url)
        # determine filename
        if outputfile:
            fname = safe_filename(outputfile)
        else:
            fname = f"{vid}_transcript.{fmt}"
        # ensure extension matches format
        if not fname.lower().endswith("." + fmt):
            fname = fname + "." + fmt
        target_path = os.path.join(outdir, fname)
        saved = save_output(vid, items, fmt, target_path)
        return {"ok": True, "path": saved, "url": url}
    except Exception as e:
        return {"ok": False, "error": str(e), "url": url, "outfile": outputfile}

def run_batch(csv_path: str, outdir: str, lang: Optional[str], make_zip: bool) -> Dict[str, Any]:
    # Ensure output directory exists
    if not os.path.exists(outdir):
        os.makedirs(outdir, exist_ok=True)

    results = []
    with open(csv_path, newline="", encoding="utf-8") as csvfile:
        reader = csv.DictReader(csvfile)

        for row in reader:
            url = row.get("URL") or row.get("Url") or row.get("url") or ""
            fmt = row.get("format") or "json"
            outfn = (
                row.get("outputfileName")
                or row.get("outputFileName")
                or row.get("output")
                or row.get("outfile")
                or None
            )

            print("Processing:", url, "->", fmt, outfn)
            res = process_row(url, fmt, outfn, outdir, lang)

            if res.get("ok"):
                print("  OK:", res["path"])
            else:
                print("  ERROR:", res.get("error"))

            results.append(res)

    zip_path = None

    if make_zip:
        zip_path = outdir.rstrip("/\\") + ".zip"

        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for root, _, files in os.walk(outdir):
                for f in files:
                    # Store files FLAT in ZIP (no directory path)
                    zf.write(os.path.join(root, f), arcname=f)

        print("Created zip:", zip_path)

        # 🔥 REMOVE THE OUTPUT DIRECTORY
        import shutil
        shutil.rmtree(outdir, ignore_errors=True)
        print("Removed directory:", outdir)

    return {"results": results, "zip": zip_path}


# ---------- PASTE THIS BLOCK near other helpers (after yt_dlp_subtitle_fallback/process_row) ----------

def extract_playlist_videos(playlist_url: str) -> List[Dict[str, str]]:
    """
    Returns list of dicts: [{'id': 'VIDEOID', 'title': 'Video Title', 'url': 'https://www.youtube.com/watch?v=VIDEOID'}, ...]
    Requires yt-dlp python package or CLI installed.
    """
    try:
        from yt_dlp import YoutubeDL
    except Exception:
        # fallback to CLI json extraction
        try:
            import subprocess, json, shlex
            cmd = ["yt-dlp", "-J", "--flat-playlist", playlist_url]
            p = subprocess.run(cmd, capture_output=True, text=True, check=True)
            info = json.loads(p.stdout)
        except Exception as e:
            raise RuntimeError(f"yt-dlp is required to fetch playlist info (python package or CLI). Error: {e}")
        entries = info.get("entries", [])
        out = []
        for e in entries:
            vid = e.get("id")
            title = e.get("title") or ""
            if vid:
                out.append({"id": vid, "title": title, "url": f"https://www.youtube.com/watch?v={vid}"})
        return out

    # using python yt_dlp
    ydl_opts = {"extract_flat": "in_playlist", "skip_download": True, "quiet": True}
    with YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(playlist_url, download=False)
    entries = info.get("entries", []) or []
    out = []
    for e in entries:
        vid = e.get("id")
        title = e.get("title") or e.get("fulltitle") or ""
        if vid:
            out.append({"id": vid, "title": title, "url": f"https://www.youtube.com/watch?v={vid}"})
    return out

def apply_output_template(template: Optional[str], index: int, video_id: str, title: str, fmt: str) -> str:
    """
    Replace placeholders in template:
      {index} -> 1-based index
      {video_id} -> video id
      {title} -> sanitized title
      {ext} -> format extension (json, srt, ...)
    If template is None -> default "{video_id}_transcript.{ext}"
    """
    safe_title = safe_filename(title) if title else ""
    if not template:
        return f"{video_id}_transcript.{fmt}"
    name = template.replace("{index}", str(index)).replace("{video_id}", video_id).replace("{title}", safe_title).replace("{ext}", fmt)
    # ensure extension matches fmt
    if not name.lower().endswith("." + fmt):
        name = name + "." + fmt
    return safe_filename(name)

# ---------- END PASTE BLOCK ----------

# --------------------------- CLI ----------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Download YouTube transcripts (single, batch, playlist)")
    parser.add_argument("url", nargs="?", help="YouTube video URL (single mode)")
    parser.add_argument("--lang", help="Language code (e.g. en)", default=None)
    parser.add_argument("--format", "-f", choices=("txt","json","srt","vtt","docx","pdf","csv"), default="json", help="output format for single mode")
    parser.add_argument("--output", "-o", help="Custom output filename for single mode", default=None)

    parser.add_argument("--batch", "-b", help="CSV file for batch mode (header: URL,format,outputfileName)", default=None)
    parser.add_argument("--outdir", help="Output directory for batch/playlist mode", default="downloads")
    parser.add_argument("--zip", action="store_true", help="Create a zip archive of outdir after batch/playlist completes")

    # PLAYLIST options
    parser.add_argument("--playlist", help="YouTube playlist URL (playlist mode). Mutually exclusive with single URL and --batch", default=None)
    parser.add_argument("--playlist-format", help="format to use for all playlist items (overrides --format)", default=None)
    parser.add_argument("--playlist-output-template", help="Filename template for playlist items. Use {index}, {video_id}, {title}, {ext}. Example: 'pl_{index}_{video_id}.{ext}'", default=None)
    parser.add_argument("--skip-existing", action="store_true", help="Skip downloading if target file already exists in outdir (useful for resuming large playlists)")

    args = parser.parse_args()

    # mode selection sanity
    mode_count = sum(bool(x) for x in (args.batch, args.playlist, args.url))
    if mode_count == 0:
        parser.print_help()
        sys.exit(1)
    if args.batch and (args.playlist or args.url):
        print("Error: --batch is mutually exclusive with single URL and --playlist")
        sys.exit(2)
    if args.playlist and args.url:
        print("Error: --playlist is mutually exclusive with single URL")
        sys.exit(2)

    # PLAYLIST MODE
    if args.playlist:
        fmt = args.playlist_format if args.playlist_format else args.format
        fmt = fmt.lower()
        if fmt not in ("txt","json","srt","vtt","docx","pdf","csv"):
            print("Unsupported playlist format:", fmt)
            sys.exit(2)
        print("Fetching playlist videos...")
        try:
            videos = extract_playlist_videos(args.playlist)
        except Exception as e:
            print("Failed to fetch playlist info:", e)
            sys.exit(3)
        if not videos:
            print("No videos found in playlist.")
            sys.exit(0)

        # ensure outdir
        outdir = args.outdir
        if not os.path.exists(outdir):
            os.makedirs(outdir, exist_ok=True)

        results = []
        for idx, v in enumerate(videos, start=1):
            video_url = v.get("url")
            vid = v.get("id")
            title = v.get("title") or ""
            outname = apply_output_template(args.playlist_output_template, idx, vid, title, fmt)
            target_path = os.path.join(outdir, outname)
            if args.skip_existing and os.path.exists(target_path):
                print(f"Skipping existing: {outname}")
                results.append({"ok": True, "path": target_path, "skipped": True, "url": video_url})
                continue
            print(f"[{idx}/{len(videos)}] Processing: {title or vid} -> {outname}")
            res = process_row(video_url, fmt, outname, outdir, args.lang)
            if res.get("ok"):
                print("  OK:", res.get("path"))
            else:
                print("  ERROR:", res.get("error"))
            results.append(res)

        zip_path = None
        if args.zip:
            zip_path = outdir.rstrip("/\\") + ".zip"
            import zipfile, shutil
            with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
                for f in os.listdir(outdir):
                    zf.write(os.path.join(outdir, f), arcname=f)
            # remove directory as you requested earlier
            shutil.rmtree(outdir, ignore_errors=True)
            print("Created zip:", zip_path)
        # summary
        ok_count = sum(1 for r in results if r.get("ok"))
        err_count = sum(1 for r in results if not r.get("ok"))
        skipped = sum(1 for r in results if r.get("skipped"))
        print(f"Playlist finished. OK: {ok_count}, ERR: {err_count}, SKIPPED: {skipped}")
        if zip_path:
            print("Zip:", zip_path)
        sys.exit(0)

    # BATCH MODE (existing behavior)
    if args.batch:
        if not os.path.exists(args.batch):
            print("Batch CSV not found:", args.batch)
            sys.exit(2)
        print("Running batch:", args.batch)
        report = run_batch(args.batch, args.outdir, args.lang, args.zip)
        ok_count = sum(1 for r in report["results"] if r.get("ok"))
        err_count = sum(1 for r in report["results"] if not r.get("ok"))
        print(f"Batch finished. OK: {ok_count}, ERR: {err_count}")
        if report.get("zip"):
            print("Zip:", report.get("zip"))
        sys.exit(0)

    # SINGLE MODE (existing)
    if args.url:
        print("Fetching transcript for single URL...")
        items = download_transcript(args.url, prefer_lang=args.lang)
        vid = extract_video_id(args.url)
        for it in items:
            it.setdefault("text", None)
            it.setdefault("start", None)
            it.setdefault("duration", None)
            it.setdefault("speaker", None)
            it.setdefault("confidence", None)
        outname = args.output if args.output else None
        outpath = save_output(vid, items, args.format, outname)
        print("Saved:", outpath)
        sys.exit(0)

if __name__ == "__main__":
    main()

